from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────
title = doc.add_heading("NTT.v Architecture Document", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("NTT / INTT Engine — Verilog Hardware Implementation").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── 1. Overview ────────────────────────────────────────────────────────────
doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "The ntt.v module implements an iterative Number Theoretic Transform (NTT) "
    "and its inverse (INTT) in hardware. It supports both the Cooley–Tukey (CT) "
    "forward transform and the Gentleman–Sande (GS) inverse transform, sharing "
    "the same datapath and memory, controlled by the inverse flag."
)

doc.add_paragraph(
    "The design is parameterised so it works for any polynomial degree N = 2^LOGN "
    "and any prime modulus q of up to Q_WIDTH bits, both supplied at runtime."
)

# ── 2. Parameters ──────────────────────────────────────────────────────────
doc.add_heading("2. Module Parameters", level=1)

param_table = doc.add_table(rows=1, cols=3)
param_table.style = "Table Grid"
hdr = param_table.rows[0].cells
hdr[0].text = "Parameter"
hdr[1].text = "Default"
hdr[2].text = "Description"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

params = [
    ("LOGN",    "13",  "log2(N) — polynomial degree N = 2^LOGN"),
    ("Q_WIDTH", "40",  "Bit-width of the prime modulus q"),
]
for name, default, desc in params:
    row = param_table.add_row().cells
    row[0].text = name
    row[1].text = default
    row[2].text = desc

doc.add_paragraph()

# ── 3. Port List ───────────────────────────────────────────────────────────
doc.add_heading("3. Port List", level=1)

port_table = doc.add_table(rows=1, cols=4)
port_table.style = "Table Grid"
hdr = port_table.rows[0].cells
for i, h in enumerate(["Port", "Direction", "Width", "Description"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True

ports = [
    ("clk",           "input",  "1",          "System clock"),
    ("rst_n",         "input",  "1",          "Active-low synchronous reset"),
    ("q",             "input",  "Q_WIDTH",    "Prime modulus (runtime)"),
    ("n_inv",         "input",  "Q_WIDTH",    "N^{-1} mod q — used in INTT scaling"),
    ("coeff_wr_en",   "input",  "1",          "Enable write to coefficient RAM"),
    ("coeff_wr_addr", "input",  "LOGN",       "Address for coefficient write"),
    ("coeff_wr_data", "input",  "Q_WIDTH",    "Data for coefficient write"),
    ("tw_wr_en",      "input",  "1",          "Enable write to twiddle ROM"),
    ("tw_wr_addr",    "input",  "LOGN+1",     "Addr [0,N) = forward; [N,2N) = inverse"),
    ("tw_wr_data",    "input",  "Q_WIDTH",    "Twiddle factor value"),
    ("start",         "input",  "1",          "Pulse high to begin transform"),
    ("inverse",       "input",  "1",          "0 = NTT (forward), 1 = INTT (inverse)"),
    ("rd_addr",       "input",  "LOGN",       "Address to read result coefficient"),
    ("rd_data",       "output", "Q_WIDTH",    "Result coefficient (1-cycle latency)"),
    ("done",          "output", "1",          "High for 1 cycle when transform complete"),
]
for p in ports:
    row = port_table.add_row().cells
    for i, v in enumerate(p):
        row[i].text = v

doc.add_paragraph()

# ── 4. Internal Architecture ───────────────────────────────────────────────
doc.add_heading("4. Internal Architecture", level=1)

doc.add_heading("4.1 Memory", level=2)
doc.add_paragraph(
    "Two on-chip memories are used:"
)
mem_items = [
    "coeff[0..N-1]  — Coefficient RAM (Q_WIDTH bits wide). Loaded before start via coeff_wr_en. "
    "In-place butterfly results overwrite this RAM during the transform.",
    "tw[0..2N-1]    — Twiddle ROM (Q_WIDTH bits wide). First half [0..N-1] holds forward NTT "
    "twiddle factors; second half [N..2N-1] holds inverse twiddle factors. "
    "Loaded before start via tw_wr_en.",
]
for item in mem_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.2 Address Generator", level=2)
doc.add_paragraph(
    "Combinational logic computes the two butterfly operand addresses (ua, va) and "
    "the twiddle index (tw_idx) from the current stage and butterfly counter k."
)

addr_table = doc.add_table(rows=1, cols=3)
addr_table.style = "Table Grid"
hdr = addr_table.rows[0].cells
for i, h in enumerate(["Signal", "NTT (CT)", "INTT (GS)"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True

addr_rows = [
    ("ts (half-size exponent)", "LOGN-1-stage  (t shrinks)", "stage         (t grows)"),
    ("ms (group exponent)",     "stage         (m grows)",   "LOGN-1-stage  (m shrinks)"),
    ("t_val",                   "1 << ts",                   "1 << ts"),
    ("m_val",                   "1 << ms",                   "1 << ms"),
    ("ua",                      "(grp << ts+1) | off",       "same formula"),
    ("va",                      "ua | t_val",                "same formula"),
    ("tw_idx",                  "m_val + grp",               "N + m_val + grp"),
]
for r in addr_rows:
    row = addr_table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

doc.add_paragraph()

doc.add_heading("4.3 Butterfly Unit", level=2)
doc.add_paragraph(
    "The butterfly is fully combinational, operating on three latched registers "
    "u_r, v_r, w_r (loaded in ST_READ). Two variants are implemented:"
)

bf_table = doc.add_table(rows=1, cols=3)
bf_table.style = "Table Grid"
hdr = bf_table.rows[0].cells
for i, h in enumerate(["Mode", "u' (upper)", "v' (lower)"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True

bf_rows = [
    ("Cooley–Tukey NTT",      "u + v·w  mod q", "u − v·w  mod q"),
    ("Gentleman–Sande INTT",  "u + v    mod q", "(u − v)·w  mod q"),
]
for r in bf_rows:
    row = bf_table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

doc.add_paragraph()
doc.add_paragraph(
    "Modular arithmetic functions (mod_add, mod_sub, mod_mul) are implemented as "
    "Verilog automatic functions. mod_mul uses the Verilog % operator — behavioural "
    "only (Verilator simulation). For FPGA/ASIC synthesis, replace mod_mul with a "
    "Barrett or Montgomery multiplier."
)

doc.add_heading("4.4 FSM", level=2)
doc.add_paragraph("The control FSM has five states:")

fsm_table = doc.add_table(rows=1, cols=3)
fsm_table.style = "Table Grid"
hdr = fsm_table.rows[0].cells
for i, h in enumerate(["State", "Encoding", "Description"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True

fsm_rows = [
    ("ST_IDLE",  "3'd0", "Wait for start pulse. Latch inverse flag."),
    ("ST_READ",  "3'd1", "Load u_r = coeff[ua], v_r = coeff[va], w_r = tw[tw_idx]."),
    ("ST_WRITE", "3'd2", "Write butterfly outputs back to coeff[ua] and coeff[va]. Advance k and stage."),
    ("ST_SCALE", "3'd3", "INTT only: multiply every coefficient by n_inv mod q (N sequential cycles)."),
    ("ST_DONE",  "3'd4", "Assert done=1 for one cycle, then return to ST_IDLE."),
]
for r in fsm_rows:
    row = fsm_table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

doc.add_paragraph()

# ── 5. FSM Transition Diagram ──────────────────────────────────────────────
doc.add_heading("5. FSM State Transition", level=1)
doc.add_paragraph("State transition sequence:")

transitions = [
    "ST_IDLE  ──(start=1)──►  ST_READ",
    "ST_READ  ──(always)──►   ST_WRITE",
    "ST_WRITE ──(k < N/2-1)──► ST_READ        (next butterfly)",
    "ST_WRITE ──(k==N/2-1, stage < LOGN-1)──► ST_READ   (next stage)",
    "ST_WRITE ──(k==N/2-1, stage==LOGN-1, NTT)──►  ST_DONE",
    "ST_WRITE ──(k==N/2-1, stage==LOGN-1, INTT)──► ST_SCALE",
    "ST_SCALE ──(sc_idx == N)──► ST_DONE",
    "ST_DONE  ──(always)──►   ST_IDLE",
]
for t in transitions:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph()

# ── 6. Timing ─────────────────────────────────────────────────────────────
doc.add_heading("6. Timing and Performance", level=1)

timing_table = doc.add_table(rows=1, cols=2)
timing_table.style = "Table Grid"
hdr = timing_table.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "Value"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

timing_rows = [
    ("Cycles per butterfly",       "2  (1 READ + 1 WRITE)"),
    ("Butterflies per stage",       "N/2"),
    ("Total NTT stages",           "LOGN"),
    ("NTT total cycles",           "LOGN × N/2 × 2  =  N × LOGN"),
    ("INTT extra scaling cycles",  "N  (one per coefficient)"),
    ("INTT total cycles",          "N × LOGN + N  =  N × (LOGN + 1)"),
    ("Read latency (rd_data)",     "1 clock cycle"),
]
for r in timing_rows:
    row = timing_table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]

doc.add_paragraph()

# ── 7. Usage Flow ─────────────────────────────────────────────────────────
doc.add_heading("7. Usage Flow", level=1)
steps = [
    "Load twiddle factors: write forward table to tw[0..N-1] and inverse table to tw[N..2N-1] via tw_wr_en.",
    "Load input polynomial: write N coefficients to coeff[0..N-1] via coeff_wr_en.",
    "Set q (prime modulus) and n_inv (N^{-1} mod q for INTT).",
    "Set inverse = 0 (NTT) or 1 (INTT).",
    "Pulse start = 1 for one clock cycle.",
    "Wait until done = 1 (one-cycle pulse).",
    "Read result coefficients via rd_addr / rd_data (one coefficient per cycle, 1-cycle latency).",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"Step {i}: {s}", style="List Number")

doc.add_paragraph()

# ── 8. Implementation Notes ────────────────────────────────────────────────
doc.add_heading("8. Implementation Notes", level=1)
notes = [
    "mod_mul uses Verilog % — this is behavioural and works only in simulation (Verilator). "
    "Replace with a Barrett or Montgomery multiplier for FPGA/ASIC synthesis.",
    "The same coefficient RAM is used for both input and output (in-place transform).",
    "Twiddle factors must be pre-computed in bit-reversed order by the host (Python side) "
    "before loading. See _ntt.py _precompute_tables().",
    "The design supports runtime-configurable modulus q — not hardcoded — so it works "
    "for any NTT-friendly prime (p ≡ 1 mod 2N).",
    "NTT and INTT share the same butterfly hardware; the inverse flag selects CT vs GS equations.",
]
for n in notes:
    doc.add_paragraph(n, style="List Bullet")

# ── Save ───────────────────────────────────────────────────────────────────
output_path = "/Users/mac/Desktop/NTT_Architecture.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
